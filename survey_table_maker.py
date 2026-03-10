import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# 1. THEME & PRESETS
# ==========================================
SURVEY_THEME = {
    'borders': {
        'table_outline': {'sz': 10, 'color': '000000', 'val': 'single'},
        'row_divider':   {'sz': 4, 'color': 'A6A6A6', 'val': 'single'},
        # 'col_divider':   {'sz': 2, 'color': 'F2F2F2', 'val': 'single'},
        'col_divider':   {'val': 'nil'},
        'header_bottom': {'sz': 10, 'color': '000000', 'val': 'single'},
    },
    'font_face': 'Arial',
    'header_fill': 'F2F2F2',
    'banding_fill': 'F9F9F9',
}

# ==========================================
# 2. XML HELPERS
# ==========================================

def set_text_rotation(cell, direction="btLr"):
    tcPr = cell._tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)
    tcPr.append(textDirection)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.insert(0, shd)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side, params in kwargs.items():
        tag = f'w:{side}'
        el = OxmlElement(tag)
        for attr, value in params.items():
            el.set(qn(f'w:{attr}'), str(value))
        tcBorders.append(el)

# ==========================================
# 3. DYNAMIC ENGINE (WITH MAX HEIGHT)
# ==========================================

def dataframe_to_word_table(doc, df, 
                            theme, 
                            rotate_headers=False, 
                            col1_min=2.0, 
                            col1_max=4.5, 
                            data_col_min=0.5, 
                            data_col_max=1.5,
                            header_max_height=1.5, 
                            data_row_max_height=0.6):
    
    rows_count, cols_count = df.shape
    table = doc.add_table(rows=rows_count + 1, cols=cols_count)
    table.autofit = False 
    
    # --- DYNAMIC WIDTHS ---
    if rotate_headers:
        final_data_width = data_col_min 
    else:
        max_header_word = max([len(str(c)) for c in df.columns[1:]])
        final_data_width = min(data_col_max, max(data_col_min, max_header_word * 0.12))

    MAX_PAGE_WIDTH = 6.25
    total_data_width = (cols_count - 1) * final_data_width
    remaining_space = MAX_PAGE_WIDTH - total_data_width
    
    longest_col1_val = max([len(str(x)) for x in df.iloc[:, 0]])
    est_col1_needed = (longest_col1_val * 0.08) + 0.2
    final_col1_width = min(col1_max, max(col1_min, min(remaining_space, est_col1_needed)))

    # Apply Widths
    for r in range(rows_count + 1):
        for c in range(cols_count):
            table.rows[r].cells[c].width = Inches(final_col1_width if c == 0 else final_data_width)

    # --- POPULATE & STYLE ---
    for r_idx in range(rows_count + 1):
        row = table.rows[r_idx]
        
        # --- NEW MAX HEIGHT LOGIC ---
        if r_idx == 0:
            if rotate_headers:
                max_header_str = max([len(str(c)) for c in df.columns[1:]])
                calc_h = (max_header_str * 0.07) + 0.1
            else:
                calc_h = 0.35
            
            # If calculated height exceeds max, cap it and switch to EXACTLY
            if calc_h > header_max_height:
                row.height = Inches(header_max_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            else:
                row.height = Inches(calc_h)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        else:
            # For data rows, use AT_LEAST with a 0.25 minimum.
            # To cap data rows, you'd use EXACTLY, but be careful of clipping long text.
            row.height = Inches(0.25)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for c_idx in range(cols_count):
            cell = row.cells[c_idx]
            
            if r_idx == 0:
                cell.text = str(df.columns[c_idx])
                if theme.get('header_fill'): set_cell_shading(cell, theme['header_fill'])
                if rotate_headers and c_idx > 0:
                    set_text_rotation(cell)
                    h_align, v_align = WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.CENTER
                else:
                    h_align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    v_align = WD_ALIGN_VERTICAL.BOTTOM
            else:
                cell.text = str(df.values[r_idx-1, c_idx])
                if r_idx % 2 == 0: set_cell_shading(cell, theme['banding_fill'])
                h_align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                v_align = WD_ALIGN_VERTICAL.CENTER

            # Fix Attribute Error by accessing the paragraph object correctly
            cell.vertical_alignment = v_align
            p = cell.paragraphs[0] 
            p.alignment = h_align
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = theme['font_face']
            run.font.size = Pt(9 if rotate_headers and r_idx == 0 and c_idx > 0 else 10)
            if r_idx == 0: run.font.bold = True

            # Borders
            b = theme['borders']
            set_cell_border(cell, **{
                'top': b['table_outline'] if r_idx == 0 else b['row_divider'],
                'left': b['table_outline'] if c_idx == 0 else b['col_divider'],
                'right': b['table_outline'] if c_idx == cols_count - 1 else b['col_divider'],
                'bottom': b['header_bottom'] if r_idx == 0 else (b['table_outline'] if r_idx == rows_count else b['row_divider'])
            })
    return table

# ==========================================
# 4. EXECUTION
# ==========================================


num_rows = 4

doc = Document()

df_likert = pd.DataFrame({
    'Survey Question': ['Comfort', 
                        'Durability',
                        'Water Resistance', 
                        'Ease of Use'],
    'Very Dissatisfied': np.random.randint(low=1, high=10, size=num_rows), 
    'Moderately Dissatisfied':np.random.randint(low=1, high=10, size=num_rows), 
    'Slightly Dissatisfied':np.random.randint(low=1, high=10, size=num_rows), 
    'Neither': np.random.randint(low=1, high=10, size=num_rows), 
    'Slightly Satisfied':np.random.randint(low=1, high=10, size=num_rows), 
    'Moderately Satisfied': np.random.randint(low=1, high=10, size=num_rows), 
    'Very Satisfied': np.random.randint(low=1, high=10, size=num_rows)
})

dataframe_to_word_table(doc, df_likert, SURVEY_THEME, rotate_headers=True, header_max_height=1.2)
doc.add_paragraph()

dataframe_to_word_table(doc, df_likert, SURVEY_THEME, rotate_headers=True, header_max_height=.6, col1_max=3)
doc.add_paragraph()


dataframe_to_word_table(doc, df_likert, SURVEY_THEME, rotate_headers=True, col1_max=3, data_col_max=.8)
doc.add_paragraph()

doc.save('Survey_Report.docx')
